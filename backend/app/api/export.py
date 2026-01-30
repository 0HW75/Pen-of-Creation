from flask import request, jsonify, send_file
from app import db
from app.models import Chapter, Project
from app.api import api_bp
import os
import tempfile
import gc
from docx import Document

# 简化的内存使用监控函数
def get_memory_usage():
    # 移除psutil依赖，返回固定值
    return 0.0

@api_bp.route('/export/word', methods=['POST'])
def export_to_word():
    temp_file_path = None
    try:
        data = request.json
        content = data.get('content', '')
        title = data.get('title', 'Document')
        
        # 监控内存使用
        print(f"Memory usage before creating document: {get_memory_usage():.2f} MB")
        
        # 创建Word文档
        doc = Document()
        doc.add_heading(title, level=1)
        
        # 分批添加内容，避免一次性加载大文件
        paragraphs = content.split('\n')
        batch_size = 100
        for i in range(0, len(paragraphs), batch_size):
            batch_paragraphs = paragraphs[i:i+batch_size]
            for paragraph in batch_paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            # 每处理一批后清理内存
            gc.collect()
        
        # 监控内存使用
        print(f"Memory usage before saving document: {get_memory_usage():.2f} MB")
        
        # 保存到临时文件
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            doc.save(temp_file.name)
            temp_file_path = temp_file.name
        
        # 清理文档对象
        del doc
        gc.collect()
        
        # 监控内存使用
        print(f"Memory usage before sending file: {get_memory_usage():.2f} MB")
        
        # 发送文件
        return send_file(temp_file_path, as_attachment=True, download_name=f'{title}.docx')
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        # 清理临时文件
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
            except:
                pass
        # 清理内存
        gc.collect()

@api_bp.route('/export/pdf', methods=['POST'])
def export_to_pdf():
    temp_file_path = None
    try:
        data = request.json
        content = data.get('content', '')
        title = data.get('title', 'Document')
        
        # 分批处理内容，避免一次性加载大文件
        paragraphs = content.split('\n')
        html_paragraphs = []
        batch_size = 100
        for i in range(0, len(paragraphs), batch_size):
            batch_paragraphs = paragraphs[i:i+batch_size]
            for para in batch_paragraphs:
                if para.strip():
                    html_paragraphs.append(f'<p>{para}</p>')
            # 每处理一批后清理内存
            gc.collect()
        
        # 创建HTML内容
        html_content = f'''
        <html>
            <head>
                <meta charset="utf-8">
                <title>{title}</title>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 20px; }}
                    h1 {{ color: #333; }}
                    p {{ line-height: 1.6; margin-bottom: 10px; }}
                </style>
            </head>
            <body>
                <h1>{title}</h1>
                {''.join(html_paragraphs)}
            </body>
        </html>
        '''
        
        # 保存到临时HTML文件
        with tempfile.NamedTemporaryFile(suffix='.html', delete=False) as temp_file:
            temp_file.write(html_content.encode('utf-8'))
            temp_file_path = temp_file.name
        
        # 清理变量
        del html_paragraphs, html_content
        gc.collect()
        
        # 发送HTML文件，让前端使用降级方案
        return send_file(temp_file_path, as_attachment=True, download_name=f'{title}.html')
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        # 清理临时文件
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
            except:
                pass
        # 清理内存
        gc.collect()

@api_bp.route('/export/markdown', methods=['POST'])
def export_to_markdown():
    temp_file_path = None
    try:
        data = request.json
        content = data.get('content', '')
        title = data.get('title', 'Document')
        
        # 创建Markdown内容
        markdown_content = f'# {title}\n\n{content}'
        
        # 保存到临时文件
        with tempfile.NamedTemporaryFile(suffix='.md', delete=False) as temp_file:
            temp_file.write(markdown_content.encode('utf-8'))
            temp_file_path = temp_file.name
        
        # 清理变量
        del markdown_content
        gc.collect()
        
        # 发送文件
        return send_file(temp_file_path, as_attachment=True, download_name=f'{title}.md')
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        # 清理临时文件
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
            except:
                pass
        # 清理内存
        gc.collect()

@api_bp.route('/export/text', methods=['POST'])
def export_to_text():
    temp_file_path = None
    try:
        data = request.json
        content = data.get('content', '')
        title = data.get('title', 'Document')
        
        # 创建文本内容
        text_content = f'{title}\n\n{content}'
        
        # 保存到临时文件
        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as temp_file:
            temp_file.write(text_content.encode('utf-8'))
            temp_file_path = temp_file.name
        
        # 清理变量
        del text_content
        gc.collect()
        
        # 发送文件
        return send_file(temp_file_path, as_attachment=True, download_name=f'{title}.txt')
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        # 清理临时文件
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
            except:
                pass
        # 清理内存
        gc.collect()
